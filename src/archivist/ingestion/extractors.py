"""File extraction and text processing for Archivist.

Supports:
- Plain text and code files (.txt, .py, .js, .ts, .java, .c, .cpp, etc.)
- PDF documents (.pdf) via pypdf
- Word documents (.docx) via python-docx

All extracted text is normalized for consistent vectorization.
"""

from __future__ import annotations

import csv
import hashlib
import io
import json
import re
from pathlib import Path
from typing import Iterator

import pypdf
from docx import Document
from openpyxl import load_workbook
import xlrd


# Supported file extensions for ingestion
SUPPORTED_EXTENSIONS = {
    # Plain text / code
    ".txt", ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".c", ".cpp", ".h",
    ".hpp", ".cc", ".cxx", ".hh", ".hxx",
    ".cs", ".rb", ".go", ".rs", ".php", ".html", ".htm", ".css", ".scss",
    ".json", ".yaml", ".yml", ".toml", ".xml", ".ini", ".cfg", ".conf",
    ".md", ".rst", ".csv", ".tsv", ".sql", ".sh", ".bash", ".zsh", ".bat",
    ".ps1", ".dockerfile", ".makefile", ".gitignore", ".env", ".log",
    # Documents
    ".pdf", ".docx",
    # Tabular data
    ".xls", ".xlsx", ".jsonl",
}

# Extensions that support line-based chunking (code/text files)
_CODE_EXTENSIONS = {
    ".c", ".cpp", ".h", ".hpp", ".cc", ".cxx", ".hh", ".hxx",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".cs", ".rb",
    ".go", ".rs", ".php", ".swift", ".kt", ".scala", ".r",
    ".sh", ".bash", ".zsh", ".bat", ".ps1",
    ".md", ".rst", ".txt", ".log",
    ".xml", ".html", ".htm", ".css", ".scss",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".sql", ".proto", ".cmake",
}

# Chunking thresholds
_CHUNK_THRESHOLD_BYTES = 10 * 1024 * 1024  # 10 MB
_CHUNK_THRESHOLD_PAGES = 100


class UnsupportedFileType(Exception):
    """Raised when attempting to extract text from an unsupported file type."""
    pass


def sha256_file(path: Path) -> str:
    """Compute SHA-256 hash of file contents.

    Args:
        path: Path to file to hash.

    Returns:
        Hex digest string of the file's SHA-256 hash.
    """
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def normalize_text(text: str) -> str:
    """Normalize text for vectorization.

    Processes text by:
    1. Converting to lowercase
    2. Stripping control characters (except newlines)
    3. Collapsing all whitespace (including newlines) to single spaces

    Args:
        text: Raw text to normalize.

    Returns:
        Normalized text string.
    """
    text = text.lower()
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_for_display(text: str) -> str:
    """Normalize text for display while preserving line structure and case.

    Strips control characters, collapses multiple spaces to one within lines,
    preserves newlines and leading indentation for line-numbered snippet display.

    Args:
        text: Raw text to normalize.

    Returns:
        Normalized text with preserved line breaks, indentation, and case.
    """
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", text)
    text = text.expandtabs(4)
    # Process each line: preserve leading spaces, collapse internal spaces
    lines = text.split("\n")
    result = []
    for line in lines:
        # Match leading whitespace and the rest
        m = re.match(r"^(\s*)(.*)", line)
        if m:
            leading = m.group(1)
            rest = re.sub(r" {2,}", " ", m.group(2))
            result.append(leading + rest)
    text = "\n".join(result)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_txt(path: Path) -> str:
    """Extract text from a plain text or code file.

    Args:
        path: Path to text file.

    Returns:
        File contents as string.
    """
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_pdf(path: Path) -> str:
    """Extract text from a PDF document.

    Args:
        path: Path to .pdf file.

    Returns:
        Concatenated text from all pages.
    """
    reader = pypdf.PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def extract_docx(path: Path) -> str:
    """Extract text from a Word document.

    Args:
        path: Path to .docx file.

    Returns:
        Concatenated text from all paragraphs.
    """
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text)


def _detect_csv_delimiter(sample: str) -> str:
    """Detect CSV delimiter from a sample of the file.

    Args:
        sample: First ~8KB of file content.

    Returns:
        Detected delimiter character.
    """
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
        return dialect.delimiter
    except csv.Error:
        return ","


def _read_file_with_fallback(path: Path) -> str:
    """Read a text file trying multiple encodings.

    Args:
        path: Path to file.

    Returns:
        File contents as string.
    """
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _flatten_csv_row(headers: list[str], row: list[str]) -> str:
    """Flatten a CSV row into searchable text.

    Args:
        headers: Column header names.
        row: Row values.

    Returns:
        Flattened string like "col1: val1 | col2: val2".
    """
    parts = []
    for h, v in zip(headers, row):
        h = h.strip()
        v = v.strip() if v else ""
        if h and v:
            parts.append(f"{h}: {v}")
        elif v:
            parts.append(v)
    return " | ".join(parts)


def extract_csv(path: Path) -> str:
    """Extract text from a CSV or TSV file with column-aware flattening.

    Auto-detects delimiter (comma, tab, semicolon, pipe). Flattens each row
    as "header: value | header: value" for optimal searchability.

    Args:
        path: Path to .csv or .tsv file.

    Returns:
        Flattened text with headers and rows.
    """
    raw = _read_file_with_fallback(path)
    sample = raw[:8192]
    delimiter = _detect_csv_delimiter(sample)

    reader = csv.reader(io.StringIO(raw), delimiter=delimiter)
    rows = list(reader)
    if not rows:
        return ""

    # Filter out empty rows
    rows = [r for r in rows if any(cell.strip() for cell in r)]
    if not rows:
        return ""

    headers = rows[0]
    lines = [" | ".join(h.strip() for h in headers if h.strip())]

    for row in rows[1:]:
        flattened = _flatten_csv_row(headers, row)
        if flattened:
            lines.append(flattened)

    return "\n".join(lines)


def extract_excel(path: Path) -> str:
    """Extract text from an Excel file (.xls or .xlsx) with column-aware flattening.

    Processes each sheet separately. Flattens rows as "header: value | header: value".

    Args:
        path: Path to .xls or .xlsx file.

    Returns:
        Flattened text with sheet names, headers, and rows.
    """
    ext = path.suffix.lower()
    lines: list[str] = []

    if ext == ".xlsx":
        wb = load_workbook(str(path), read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            lines.append(f"[Sheet: {sheet_name}]")
            headers = [str(c) if c is not None else "" for c in rows[0]]
            lines.append(" | ".join(h for h in headers if h))
            for row in rows[1:]:
                cells = [str(c) if c is not None else "" for c in row]
                flattened = _flatten_csv_row(headers, cells)
                if flattened:
                    lines.append(flattened)
        wb.close()
    elif ext == ".xls":
        wb = xlrd.open_workbook(str(path))
        for sheet_idx in range(wb.nsheets):
            ws = wb.sheet_by_index(sheet_idx)
            if ws.nrows == 0:
                continue
            lines.append(f"[Sheet: {ws.name}]")
            headers = [str(ws.cell_value(0, c)) for c in range(ws.ncols)]
            lines.append(" | ".join(h for h in headers if h))
            for row_idx in range(1, ws.nrows):
                cells = [str(ws.cell_value(row_idx, c)) for c in range(ws.ncols)]
                flattened = _flatten_csv_row(headers, cells)
                if flattened:
                    lines.append(flattened)

    return "\n".join(lines)


def _flatten_json_obj(obj: dict, prefix: str = "") -> list[str]:
    """Flatten a JSON object into dot-notation key-value pairs.

    Args:
        obj: Dictionary to flatten.
        prefix: Current key prefix for nested objects.

    Returns:
        List of "key: value" strings.
    """
    parts: list[str] = []
    for k, v in obj.items():
        key = f"{prefix}.{k}" if prefix else k
        if isinstance(v, dict):
            parts.extend(_flatten_json_obj(v, key))
        elif isinstance(v, list):
            parts.append(f"{key}: {json.dumps(v, ensure_ascii=False)}")
        elif v is not None:
            parts.append(f"{key}: {str(v)}")
    return parts


def extract_jsonl(path: Path) -> str:
    """Extract text from a JSON Lines file.

    Each line is parsed as a JSON object and flattened with dot-notation
    for nested keys. Line numbers are included for context.

    Args:
        path: Path to .jsonl file.

    Returns:
        Flattened text with line-numbered entries.
    """
    raw = _read_file_with_fallback(path)
    lines: list[str] = []

    for i, line in enumerate(raw.splitlines(), 1):
        line = line.strip()
        if not line:
            continue
        try:
            obj = json.loads(line)
            if isinstance(obj, dict):
                flattened = _flatten_json_obj(obj)
                lines.append(f"L{i}: {' | '.join(flattened)}")
            elif isinstance(obj, list):
                lines.append(f"L{i}: {json.dumps(obj, ensure_ascii=False)}")
            else:
                lines.append(f"L{i}: {str(obj)}")
        except json.JSONDecodeError:
            lines.append(f"L{i}: {line}")

    return "\n".join(lines)


def extract_text(path: Path) -> str:
    """Extract text from a file based on its extension.

    Args:
        path: Path to file.

    Returns:
        Extracted text content.

    Raises:
        UnsupportedFileType: If file extension is not supported.
    """
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext in (".csv", ".tsv"):
        return extract_csv(path)
    if ext in (".xls", ".xlsx"):
        return extract_excel(path)
    if ext == ".jsonl":
        return extract_jsonl(path)
    # All text and code files — read as plain text
    if ext in SUPPORTED_EXTENSIONS:
        return extract_txt(path)
    raise UnsupportedFileType(f"Unsupported file type: {ext}")


def should_chunk(path: Path, text: str) -> bool:
    """Determine if a document should be split into chunks.

    Args:
        path: Path to file.
        text: Extracted text content.

    Returns:
        True if document exceeds size or page thresholds.
    """
    size = path.stat().st_size
    if size > _CHUNK_THRESHOLD_BYTES:
        return True
    ext = path.suffix.lower()
    if ext == ".pdf":
        try:
            page_count = len(pypdf.PdfReader(str(path)).pages)
            if page_count > _CHUNK_THRESHOLD_PAGES:
                return True
        except Exception:
            pass
    if ext in (".csv", ".tsv"):
        line_count = text.count("\n") + 1
        if line_count > 5000:
            return True
    if ext == ".jsonl":
        line_count = text.count("\n") + 1
        if line_count > 10000:
            return True
    if ext in _CODE_EXTENSIONS:
        line_count = text.count("\n") + 1
        if line_count > 500:
            return True
    return False


def chunk_pdf_by_page(path: Path) -> list[str]:
    """Split PDF into per-page text chunks.

    Args:
        path: Path to PDF file.

    Returns:
        List of text strings, one per page.
    """
    reader = pypdf.PdfReader(str(path))
    return [p.extract_text() or "" for p in reader.pages]


def chunk_docx_by_section(path: Path) -> list[str]:
    """Split DOCX by heading-level paragraphs.

    Args:
        path: Path to DOCX file.

    Returns:
        List of text chunks split at heading boundaries.
    """
    doc = Document(str(path))
    chunks: list[str] = []
    current: list[str] = []
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading") and current:
            chunks.append("\n".join(current))
            current = []
        current.append(para.text)
    if current:
        chunks.append("\n".join(current))
    return chunks if chunks else ["\n".join(p.text for p in doc.paragraphs if p.text)]


def chunk_text(path: Path, text: str) -> list[str]:
    """Chunk large documents by page (PDF), section (DOCX), rows (CSV), lines (JSONL),
    or lines (code/text files).

    Args:
        path: Path to file.
        text: Extracted text content.

    Returns:
        List of text chunks.
    """
    ext = path.suffix.lower()
    if ext == ".pdf":
        return chunk_pdf_by_page(path)
    if ext == ".docx":
        return chunk_docx_by_section(path)
    if ext in (".csv", ".tsv"):
        return chunk_csv_by_rows(text)
    if ext == ".jsonl":
        return chunk_jsonl_by_lines(text)
    if ext in _CODE_EXTENSIONS:
        return chunk_code_by_lines(text)
    return [text]


def chunk_csv_by_rows(text: str, rows_per_chunk: int = 1000) -> list[str]:
    """Split CSV text into chunks by row count, preserving headers in each chunk.

    Args:
        text: Flattened CSV text from extract_csv().
        rows_per_chunk: Maximum rows per chunk.

    Returns:
        List of text chunks.
    """
    lines = text.split("\n")
    if not lines:
        return [text]

    header = lines[0]
    data_lines = lines[1:]
    if not data_lines:
        return [text]

    chunks: list[str] = []
    for i in range(0, len(data_lines), rows_per_chunk):
        chunk_lines = [header] + data_lines[i : i + rows_per_chunk]
        chunks.append("\n".join(chunk_lines))
    return chunks


def chunk_jsonl_by_lines(text: str, lines_per_chunk: int = 500) -> list[str]:
    """Split JSONL text into chunks by line count.

    Args:
        text: Flattened JSONL text from extract_jsonl().
        lines_per_chunk: Maximum lines per chunk.

    Returns:
        List of text chunks.
    """
    lines = text.split("\n")
    if not lines:
        return [text]

    chunks: list[str] = []
    for i in range(0, len(lines), lines_per_chunk):
        chunks.append("\n".join(lines[i : i + lines_per_chunk]))
    return chunks


def chunk_code_by_lines(text: str, lines_per_chunk: int = 1500) -> list[str]:
    """Split code/text into chunks by line count.

    Used for source code, config, markdown, and other text files
    that benefit from line-based chunking instead of content truncation.

    Args:
        text: Extracted text content.
        lines_per_chunk: Maximum lines per chunk (default 1500).

    Returns:
        List of text chunks, each up to lines_per_chunk lines.
    """
    lines = text.split("\n")
    if len(lines) <= lines_per_chunk:
        return [text]
    chunks: list[str] = []
    for i in range(0, len(lines), lines_per_chunk):
        chunks.append("\n".join(lines[i : i + lines_per_chunk]))
    return chunks


def iter_files(root: Path, recursive: bool = True) -> Iterator[Path]:
    """Yield supported files from a directory.

    Args:
        root: Root directory or file to scan.
        recursive: If True, recurse into subdirectories.

    Yields:
        Path objects for supported files.
    """
    if root.is_file():
        if _is_supported(root):
            yield root
        return
    pattern = "**/*" if recursive else "*"
    for p in root.glob(pattern):
        if p.is_file() and _is_supported(p):
            yield p


def _is_supported(path: Path) -> bool:
    """Check if a file is supported for ingestion.

    Matches by extension or by known extensionless filenames.

    Args:
        path: Path to check.

    Returns:
        True if file can be ingested.
    """
    ext = path.suffix.lower()
    if ext in SUPPORTED_EXTENSIONS:
        return True
    # Extensionless files like Makefile, Dockerfile, .gitignore
    if path.name.lower() in {"makefile", "dockerfile", "procfile"}:
        return True
    return False
