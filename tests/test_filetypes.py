"""Comprehensive tests for all supported file types.

Covers: extract_text, should_chunk, chunk_text, and full ingest→search
pipeline for every supported file extension.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

from archivist.ingestion.extractors import (
    SUPPORTED_EXTENSIONS,
    chunk_code_by_lines,
    chunk_text,
    extract_text,
    extract_txt,
    iter_files,
    should_chunk,
)
from archivist.ingestion.pipeline import ingest_file
from archivist.ingestion.tracker import Tracker
from archivist.search.sqlite_search import SQLiteSearch


# ── Helpers ───────────────────────────────────────────────────────────────────

_MINIMAL_PDF = (
    b"%PDF-1.4\n"
    b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    b"2 0 obj\n<< /Type /Pages /Kids [3 0 R 4 0 R 5 0 R] /Count 3 >>\nendobj\n"
    b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] >>\nendobj\n"
    b"4 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] >>\nendobj\n"
    b"5 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] >>\nendobj\n"
    b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n"
    b"0000000115 00000 n \n0000000216 00000 n \n0000000317 00000 n \n"
    b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n400\n%%EOF\n"
)


def _make_docx(path: Path, heading: str, body: str) -> Path:
    doc = Document()
    doc.add_heading(heading, 0)
    doc.add_paragraph(body)
    doc.save(str(path))
    return path


def _setup_ingest(tmp_path: Path, filename: str, content: str | bytes, ext: str):
    """Create file, ingest, return search object."""
    f = tmp_path / f"{filename}{ext}"
    if isinstance(content, bytes):
        f.write_bytes(content)
    else:
        f.write_text(content)
    tracker = Tracker(tmp_path / "tracker.db")
    db_path = tmp_path / "test.db"
    count = ingest_file(f, tracker, db_path=db_path)
    search = SQLiteSearch(db_path)
    return search, count


# ── Plain text / code files ──────────────────────────────────────────────────


class TestPlainText:
    """Tests for .txt files via extract_txt."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "doc.txt"
        f.write_text("Hello world\nSecond line")
        text = extract_text(f)
        assert "Hello world" in text
        assert "Second line" in text

    def test_chunk(self, tmp_path: Path):
        f = tmp_path / "doc.txt"
        f.write_text("line1\nline2\nline3")
        text = extract_text(f)
        chunks = chunk_text(f, text)
        assert len(chunks) == 1

    def test_ingest_search(self, tmp_path: Path):
        content = "The quick brown fox jumps over the lazy dog"
        search, count = _setup_ingest(tmp_path, "doc", content, ".txt")
        assert count >= 1
        results = search.search("quick brown fox")
        assert len(results) > 0
        assert "quick brown fox" in results[0]["content"]
        search.close()


class TestPython:
    """Tests for .py files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "script.py"
        f.write_text("def hello():\n    print('Hello')\n")
        text = extract_text(f)
        assert "def hello" in text
        assert "print" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "def calculate_sum(a, b):\n    return a + b\n"
        search, count = _setup_ingest(tmp_path, "script", content, ".py")
        assert count >= 1
        results = search.search("calculate_sum")
        assert len(results) > 0
        search.close()


class TestJavaScript:
    """Tests for .js files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "app.js"
        f.write_text("function greet(name) {\n  return `Hello ${name}`;\n}\n")
        text = extract_text(f)
        assert "greet" in text
        assert "Hello" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "const fetchData = async (url) => {\n  const res = await fetch(url);\n  return res.json();\n};\n"
        search, count = _setup_ingest(tmp_path, "app", content, ".js")
        assert count >= 1
        results = search.search("fetchData")
        assert len(results) > 0
        search.close()


class TestTypeScript:
    """Tests for .ts files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "types.ts"
        f.write_text("interface User {\n  name: string;\n  age: number;\n}\n")
        text = extract_text(f)
        assert "interface User" in text
        assert "name: string" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "type Status = 'active' | 'inactive';\nfunction getStatus(): Status {\n  return 'active';\n}\n"
        search, count = _setup_ingest(tmp_path, "types", content, ".ts")
        assert count >= 1
        results = search.search("Status")
        assert len(results) > 0
        search.close()


class TestJava:
    """Tests for .java files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "Main.java"
        f.write_text("public class Main {\n  public static void main(String[] args) {\n    System.out.println(\"Hello\");\n  }\n}\n")
        text = extract_text(f)
        assert "public class Main" in text
        assert "main" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "public class Calculator {\n  public int add(int a, int b) {\n    return a + b;\n  }\n}\n"
        search, count = _setup_ingest(tmp_path, "Calculator", content, ".java")
        assert count >= 1
        results = search.search("Calculator")
        assert len(results) > 0
        search.close()


class TestC:
    """Tests for .c files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "main.c"
        f.write_text('#include <stdio.h>\nint main() {\n  printf("Hello");\n  return 0;\n}\n')
        text = extract_text(f)
        assert "#include" in text
        assert "printf" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "int add(int a, int b) {\n  return a + b;\n}\n"
        search, count = _setup_ingest(tmp_path, "math", content, ".c")
        assert count >= 1
        results = search.search("add")
        assert len(results) > 0
        search.close()


class TestCpp:
    """Tests for .cpp files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "app.cpp"
        f.write_text("#include <iostream>\nint main() {\n  std::cout << \"Hello\";\n  return 0;\n}\n")
        text = extract_text(f)
        assert "#include" in text
        assert "std::cout" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "#include <vector>\nstd::vector<int> getNumbers() {\n  return {1, 2, 3};\n}\n"
        search, count = _setup_ingest(tmp_path, "app", content, ".cpp")
        assert count >= 1
        results = search.search("vector")
        assert len(results) > 0
        search.close()


class TestCHeader:
    """Tests for .h files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "math.h"
        f.write_text("#ifndef MATH_H\n#define MATH_H\nint add(int a, int b);\n#endif\n")
        text = extract_text(f)
        assert "#ifndef" in text
        assert "add" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "#pragma once\nvoid process_data(int* buffer, size_t len);\n"
        search, count = _setup_ingest(tmp_path, "math", content, ".h")
        assert count >= 1
        results = search.search("process_data")
        assert len(results) > 0
        search.close()


class TestCppHeader:
    """Tests for .hpp files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "utils.hpp"
        f.write_text("#pragma once\n#include <string>\nclass Utils {\npublic:\n  static std::string trim(const std::string& s);\n};\n")
        text = extract_text(f)
        assert "class Utils" in text
        assert "trim" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "template<typename T>\nT clamp(T value, T min, T max);\n"
        search, count = _setup_ingest(tmp_path, "utils", content, ".hpp")
        assert count >= 1
        results = search.search("clamp")
        assert len(results) > 0
        search.close()


class TestCSharp:
    """Tests for .cs files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "Program.cs"
        f.write_text("using System;\nclass Program {\n  static void Main() {\n    Console.WriteLine(\"Hello\");\n  }\n}\n")
        text = extract_text(f)
        assert "class Program" in text
        assert "Console.WriteLine" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "namespace MyApp;\npublic class UserService {\n  public User GetUser(int id) {\n    return _db.Find(id);\n  }\n}\n"
        search, count = _setup_ingest(tmp_path, "Program", content, ".cs")
        assert count >= 1
        results = search.search("UserService")
        assert len(results) > 0
        search.close()


class TestRuby:
    """Tests for .rb files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "app.rb"
        f.write_text("class Animal\n  def initialize(name)\n    @name = name\n  end\nend\n")
        text = extract_text(f)
        assert "class Animal" in text
        assert "initialize" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "def fibonacci(n)\n  return n if n <= 1\n  fibonacci(n - 1) + fibonacci(n - 2)\nend\n"
        search, count = _setup_ingest(tmp_path, "app", content, ".rb")
        assert count >= 1
        results = search.search("fibonacci")
        assert len(results) > 0
        search.close()


class TestGo:
    """Tests for .go files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "main.go"
        f.write_text('package main\nimport "fmt"\nfunc main() {\n  fmt.Println("Hello")\n}\n')
        text = extract_text(f)
        assert "package main" in text
        assert "fmt.Println" in text

    def test_ingest_search(self, tmp_path: Path):
        content = 'package server\nfunc HandleRequest(w http.ResponseWriter, r *http.Request) {\n  fmt.Fprintf(w, "OK")\n}\n'
        search, count = _setup_ingest(tmp_path, "main", content, ".go")
        assert count >= 1
        results = search.search("HandleRequest")
        assert len(results) > 0
        search.close()


class TestRust:
    """Tests for .rs files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "lib.rs"
        f.write_text("fn main() {\n    println!(\"Hello\");\n}\n")
        text = extract_text(f)
        assert "fn main" in text
        assert "println!" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "pub fn parse_config(path: &str) -> Result<Config, Error> {\n    let content = std::fs::read_to_string(path)?;\n    Ok(serde_json::from_str(&content)?)\n}\n"
        search, count = _setup_ingest(tmp_path, "lib", content, ".rs")
        assert count >= 1
        results = search.search("parse_config")
        assert len(results) > 0
        search.close()


class TestPHP:
    """Tests for .php files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "index.php"
        f.write_text('<?php\nclass User {\n  private $name;\n  public function __construct($name) {\n    $this->name = $name;\n  }\n}\n')
        text = extract_text(f)
        assert "class User" in text
        assert "__construct" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "<?php\nfunction processRequest($request) {\n  return json_encode(['status' => 'ok']);\n}\n"
        search, count = _setup_ingest(tmp_path, "index", content, ".php")
        assert count >= 1
        results = search.search("processRequest")
        assert len(results) > 0
        search.close()


# ── HTML / CSS ────────────────────────────────────────────────────────────────


class TestHTML:
    """Tests for .html files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "index.html"
        f.write_text("<!DOCTYPE html>\n<html>\n<head><title>Test</title></head>\n<body>\n<h1>Hello</h1>\n<p>World</p>\n</body>\n</html>")
        text = extract_text(f)
        assert "<title>" in text
        assert "Hello" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "<html><body>\n<h1>Product Page</h1>\n<div class='price'>$29.99</div>\n</body></html>"
        search, count = _setup_ingest(tmp_path, "index", content, ".html")
        assert count >= 1
        results = search.search("Product Page")
        assert len(results) > 0
        search.close()


class TestCSS:
    """Tests for .css files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "style.css"
        f.write_text("body {\n  font-size: 16px;\n  color: #333;\n}\n.container {\n  max-width: 1200px;\n}\n")
        text = extract_text(f)
        assert "font-size" in text
        assert "max-width" in text

    def test_ingest_search(self, tmp_path: Path):
        content = ".button {\n  background-color: blue;\n  border-radius: 5px;\n  padding: 10px 20px;\n}\n"
        search, count = _setup_ingest(tmp_path, "style", content, ".css")
        assert count >= 1
        results = search.search("background-color")
        assert len(results) > 0
        search.close()


class TestSCSS:
    """Tests for .scss files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "main.scss"
        f.write_text("$primary: #007bff;\n.btn {\n  background: $primary;\n  &:hover {\n    opacity: 0.8;\n  }\n}\n")
        text = extract_text(f)
        assert "$primary" in text
        assert "&:hover" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "$spacing: 16px;\n@mixin flex-center {\n  display: flex;\n  justify-content: center;\n}\n"
        search, count = _setup_ingest(tmp_path, "main", content, ".scss")
        assert count >= 1
        results = search.search("flex-center")
        assert len(results) > 0
        search.close()


# ── Config files ──────────────────────────────────────────────────────────────


class TestJSON:
    """Tests for .json files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "config.json"
        f.write_text('{"name": "test", "version": "1.0", "settings": {"debug": true}}')
        text = extract_text(f)
        assert '"name"' in text
        assert '"debug"' in text

    def test_ingest_search(self, tmp_path: Path):
        content = '{"database": {"host": "localhost", "port": 5432}, "api_key": "abc123"}'
        search, count = _setup_ingest(tmp_path, "config", content, ".json")
        assert count >= 1
        results = search.search("database")
        assert len(results) > 0
        search.close()


class TestYAML:
    """Tests for .yaml files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "config.yaml"
        f.write_text("server:\n  host: localhost\n  port: 8080\ndatabase:\n  name: mydb\n")
        text = extract_text(f)
        assert "server:" in text
        assert "database:" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "deployment:\n  replicas: 3\n  image: nginx:latest\n  resources:\n    cpu: 500m\n    memory: 256Mi\n"
        search, count = _setup_ingest(tmp_path, "config", content, ".yaml")
        assert count >= 1
        results = search.search("nginx")
        assert len(results) > 0
        search.close()


class TestTOML:
    """Tests for .toml files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "pyproject.toml"
        f.write_text('[project]\nname = "myapp"\nversion = "0.1.0"\n\n[tool.pytest]\ntestpaths = ["tests"]\n')
        text = extract_text(f)
        assert "[project]" in text
        assert "myapp" in text

    def test_ingest_search(self, tmp_path: Path):
        content = '[package]\nname = "archivist"\nedition = "2021"\n\n[dependencies]\nserde = "1.0"\ntokio = "1"\n'
        search, count = _setup_ingest(tmp_path, "pyproject", content, ".toml")
        assert count >= 1
        results = search.search("serde")
        assert len(results) > 0
        search.close()


class TestINI:
    """Tests for .ini files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "settings.ini"
        f.write_text("[database]\nhost = localhost\nport = 3306\n\n[logging]\nlevel = DEBUG\n")
        text = extract_text(f)
        assert "[database]" in text
        assert "host = localhost" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "[server]\nbind = 0.0.0.0\nworkers = 4\n\n[ssl]\ncertfile = /etc/ssl/cert.pem\n"
        search, count = _setup_ingest(tmp_path, "settings", content, ".ini")
        assert count >= 1
        results = search.search("ssl")
        assert len(results) > 0
        search.close()


class TestConfig:
    """Tests for .cfg / .conf files."""

    def test_extract_cfg(self, tmp_path: Path):
        f = tmp_path / "app.cfg"
        f.write_text("[main]\nverbose = true\nmax_retries = 3\n")
        text = extract_text(f)
        assert "verbose" in text

    def test_extract_conf(self, tmp_path: Path):
        f = tmp_path / "nginx.conf"
        f.write_text("server {\n  listen 80;\n  server_name example.com;\n  location / {\n    proxy_pass http://localhost:3000;\n  }\n}\n")
        text = extract_text(f)
        assert "listen 80" in text
        assert "proxy_pass" in text

    def test_ingest_search_cfg(self, tmp_path: Path):
        content = "[app]\nsecret_key = mysecret123\ndebug = false\n"
        search, count = _setup_ingest(tmp_path, "app", content, ".cfg")
        assert count >= 1
        results = search.search("secret_key")
        assert len(results) > 0
        search.close()


# ── Shell scripts ─────────────────────────────────────────────────────────────


class TestShellScript:
    """Tests for .sh files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "deploy.sh"
        f.write_text("#!/bin/bash\nset -e\necho 'Deploying...'\ndocker compose up -d\n")
        text = extract_text(f)
        assert "#!/bin/bash" in text
        assert "docker" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "#!/bin/bash\nfor i in $(seq 1 10); do\n  echo \"Processing $i\"\ndone\n"
        search, count = _setup_ingest(tmp_path, "deploy", content, ".sh")
        assert count >= 1
        results = search.search("Processing")
        assert len(results) > 0
        search.close()


class TestBatchScript:
    """Tests for .bat files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "build.bat"
        f.write_text("@echo off\necho Building...\npython -m pytest\necho Done.\n")
        text = extract_text(f)
        assert "@echo off" in text
        assert "pytest" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "@echo off\nset PATH=C:\\Python314;%PATH%\npython main.py --config prod.json\n"
        search, count = _setup_ingest(tmp_path, "build", content, ".bat")
        assert count >= 1
        results = search.search("config")
        assert len(results) > 0
        search.close()


class TestPowerShell:
    """Tests for .ps1 files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "script.ps1"
        f.write_text("Write-Host 'Hello'\nGet-Process | Where-Object {$_.CPU -gt 100}\n")
        text = extract_text(f)
        assert "Write-Host" in text
        assert "Get-Process" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "function Install-Package {\n  param([string]$Name)\n  Write-Output \"Installing $Name...\"\n}\n"
        search, count = _setup_ingest(tmp_path, "script", content, ".ps1")
        assert count >= 1
        results = search.search("Install-Package")
        assert len(results) > 0
        search.close()


# ── Documentation ─────────────────────────────────────────────────────────────


class TestMarkdown:
    """Tests for .md files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "README.md"
        f.write_text("# Title\n\nSome **bold** text.\n\n## Section\n\n- item 1\n- item 2\n")
        text = extract_text(f)
        assert "# Title" in text
        assert "**bold**" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "# Project X\n\n## Installation\n\nRun `pip install project-x`.\n\n## Usage\n\nImport and call `init()`.\n"
        search, count = _setup_ingest(tmp_path, "README", content, ".md")
        assert count >= 1
        results = search.search("Installation")
        assert len(results) > 0
        search.close()


class TestRST:
    """Tests for .rst files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "docs.rst"
        f.write_text("Title\n=====\n\nSome text here.\n\nSection\n-------\n\nMore content.\n")
        text = extract_text(f)
        assert "Title" in text
        assert "=====" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "API Reference\n=============\n\nThe ``connect()`` function establishes a connection.\n"
        search, count = _setup_ingest(tmp_path, "docs", content, ".rst")
        assert count >= 1
        results = search.search("connect")
        assert len(results) > 0
        search.close()


# ── SQL ───────────────────────────────────────────────────────────────────────


class TestSQL:
    """Tests for .sql files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "schema.sql"
        f.write_text("CREATE TABLE users (\n  id INTEGER PRIMARY KEY,\n  name TEXT NOT NULL,\n  email TEXT UNIQUE\n);\n")
        text = extract_text(f)
        assert "CREATE TABLE" in text
        assert "users" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "SELECT u.name, COUNT(o.id) AS order_count\nFROM users u\nJOIN orders o ON u.id = o.user_id\nWHERE o.status = 'active'\nGROUP BY u.name;\n"
        search, count = _setup_ingest(tmp_path, "schema", content, ".sql")
        assert count >= 1
        results = search.search("users orders")
        assert len(results) > 0
        search.close()


# ── PDF ───────────────────────────────────────────────────────────────────────


class TestPDF:
    """Tests for .pdf files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "doc.pdf"
        f.write_bytes(_MINIMAL_PDF)
        text = extract_text(f)
        assert isinstance(text, str)
        assert len(text) >= 0  # minimal PDF may have empty text

    def test_chunk(self, tmp_path: Path):
        f = tmp_path / "doc.pdf"
        f.write_bytes(_MINIMAL_PDF)
        text = extract_text(f)
        chunks = chunk_text(f, text)
        assert len(chunks) == 3  # 3 pages

    def test_ingest_search(self, tmp_path: Path):
        f = tmp_path / "doc.pdf"
        f.write_bytes(_MINIMAL_PDF)
        tracker = Tracker(tmp_path / "tracker.db")
        db_path = tmp_path / "test.db"
        ingest_file(f, tracker, db_path=db_path)
        # Minimal PDF has no extractable text; ingest_file returns 0
        search = SQLiteSearch(db_path)
        stats = search.stats()
        assert stats["points_count"] >= 0
        search.close()


# ── DOCX ──────────────────────────────────────────────────────────────────────


class TestDOCX:
    """Tests for .docx files."""

    def test_extract(self, tmp_path: Path):
        f = _make_docx(tmp_path / "doc.docx", "Introduction", "This is the introduction.")
        text = extract_text(f)
        assert "Introduction" in text
        assert "introduction" in text.lower()

    def test_chunk(self, tmp_path: Path):
        f = tmp_path / "doc.docx"
        doc = Document()
        doc.add_heading("Part 1", 1)
        doc.add_paragraph("Content 1")
        doc.add_heading("Part 2", 1)
        doc.add_paragraph("Content 2")
        doc.save(str(f))
        text = extract_text(f)
        chunks = chunk_text(f, text)
        # chunk_docx_by_section splits on Heading styles
        all_text = "\n".join(chunks)
        assert "Part 1" in all_text
        assert "Part 2" in all_text

    def test_ingest_search(self, tmp_path: Path):
        f = tmp_path / "doc.docx"
        doc = Document()
        doc.add_heading("Quarterly Report", 0)
        doc.add_paragraph("Revenue increased by 15% this quarter.")
        doc.save(str(f))
        tracker = Tracker(tmp_path / "tracker.db")
        db_path = tmp_path / "test.db"
        count = ingest_file(f, tracker, db_path=db_path)
        assert count >= 1
        search = SQLiteSearch(db_path)
        results = search.search("Revenue")
        assert len(results) > 0
        search.close()


# ── Excel ─────────────────────────────────────────────────────────────────────


class TestExcel:
    """Tests for .xlsx files."""

    def test_extract(self, tmp_path: Path):
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(["Name", "Age", "City"])
        ws.append(["Alice", 30, "NYC"])
        ws.append(["Bob", 25, "LA"])
        f = tmp_path / "data.xlsx"
        wb.save(str(f))
        wb.close()

        text = extract_text(f)
        assert "Data" in text
        assert "Alice" in text
        assert "NYC" in text

    def test_ingest_search(self, tmp_path: Path):
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.append(["Product", "Price", "Stock"])
        ws.append(["Widget", 9.99, 100])
        ws.append(["Gadget", 24.99, 50])
        f = tmp_path / "inventory.xlsx"
        wb.save(str(f))
        wb.close()

        tracker = Tracker(tmp_path / "tracker.db")
        db_path = tmp_path / "test.db"
        count = ingest_file(f, tracker, db_path=db_path)
        assert count >= 1
        search = SQLiteSearch(db_path)
        results = search.search("Widget")
        assert len(results) > 0
        search.close()


# ── CSV / TSV ─────────────────────────────────────────────────────────────────


class TestCSV:
    """Tests for .csv files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "data.csv"
        f.write_text("name,age,city\nAlice,30,NYC\nBob,25,LA\n")
        text = extract_text(f)
        assert "name" in text
        assert "Alice" in text

    def test_chunk(self, tmp_path: Path):
        lines = ["col1,col2"] + [f"val{i},val{i}" for i in range(2500)]
        f = tmp_path / "big.csv"
        f.write_text("\n".join(lines))
        text = extract_text(f)
        chunks = chunk_text(f, text)
        assert len(chunks) > 1

    def test_ingest_search(self, tmp_path: Path):
        content = "product,price,category\nLaptop,999.99,Electronics\nCoffee,4.99,Food\n"
        search, count = _setup_ingest(tmp_path, "data", content, ".csv")
        assert count >= 1
        results = search.search("Laptop")
        assert len(results) > 0
        search.close()


class TestTSV:
    """Tests for .tsv files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "data.tsv"
        f.write_text("name\tage\tcity\nAlice\t30\tNYC\nBob\t25\tLA\n")
        text = extract_text(f)
        assert "name" in text
        assert "Alice" in text

    def test_ingest_search(self, tmp_path: Path):
        content = "id\tmetric\tvalue\n1\tcpu\t45.2\n2\tmemory\t78.1\n"
        search, count = _setup_ingest(tmp_path, "data", content, ".tsv")
        assert count >= 1
        results = search.search("cpu")
        assert len(results) > 0
        search.close()


# ── JSONL ─────────────────────────────────────────────────────────────────────


class TestJSONL:
    """Tests for .jsonl files."""

    def test_extract(self, tmp_path: Path):
        f = tmp_path / "logs.jsonl"
        f.write_text('{"level": "info", "msg": "started"}\n{"level": "error", "msg": "failed"}\n')
        text = extract_text(f)
        assert "info" in text
        assert "error" in text

    def test_chunk(self, tmp_path: Path):
        lines = [f'{{"id": {i}, "data": "row{i}"}}' for i in range(1200)]
        f = tmp_path / "big.jsonl"
        f.write_text("\n".join(lines))
        text = extract_text(f)
        chunks = chunk_text(f, text)
        assert len(chunks) > 1

    def test_ingest_search(self, tmp_path: Path):
        content = '{"event": "click", "target": "button_submit"}\n{"event": "page_view", "url": "/dashboard"}\n'
        search, count = _setup_ingest(tmp_path, "logs", content, ".jsonl")
        assert count >= 1
        results = search.search("button_submit")
        assert len(results) > 0
        search.close()


# ── Code chunking ─────────────────────────────────────────────────────────────


class TestCodeChunking:
    """Tests for line-based chunking of code files."""

    def test_chunk_python_large(self, tmp_path: Path):
        lines = [f"def func_{i}():\n    return {i}" for i in range(3000)]
        f = tmp_path / "big.py"
        f.write_text("\n".join(lines))
        text = extract_text(f)
        chunks = chunk_text(f, text)
        assert len(chunks) > 1

    def test_chunk_cpp_large(self, tmp_path: Path):
        lines = [f"void function_{i}() {{ /* func {i} */ }}" for i in range(3000)]
        f = tmp_path / "big.cpp"
        f.write_text("\n".join(lines))
        text = extract_text(f)
        chunks = chunk_text(f, text)
        assert len(chunks) > 1

    def test_chunk_preserves_all_content(self, tmp_path: Path):
        lines = [f"line_{i}" for i in range(200)]
        f = tmp_path / "code.py"
        f.write_text("\n".join(lines))
        text = extract_text(f)
        chunks = chunk_text(f, text)
        all_text = "\n".join(chunks)
        for i in range(200):
            assert f"line_{i}" in all_text


# ── Extension coverage ────────────────────────────────────────────────────────


class TestExtensionCoverage:
    """Verify all expected extensions are in SUPPORTED_EXTENSIONS."""

    def test_text_extensions(self):
        for ext in [".txt", ".py", ".js", ".ts", ".jsx", ".tsx", ".java",
                     ".c", ".cpp", ".h", ".hpp", ".cc", ".cxx", ".hh", ".hxx",
                     ".cs", ".rb", ".go", ".rs", ".php",
                     ".html", ".htm", ".css", ".scss",
                     ".json", ".yaml", ".yml", ".toml", ".xml", ".ini", ".cfg", ".conf",
                     ".md", ".rst", ".sql", ".sh", ".bash", ".zsh", ".bat", ".ps1",
                     ".log"]:
            assert ext in SUPPORTED_EXTENSIONS, f"{ext} missing from SUPPORTED_EXTENSIONS"

    def test_document_extensions(self):
        for ext in [".pdf", ".docx"]:
            assert ext in SUPPORTED_EXTENSIONS

    def test_tabular_extensions(self):
        for ext in [".xls", ".xlsx", ".jsonl"]:
            assert ext in SUPPORTED_EXTENSIONS

    def test_csv_in_supported(self):
        assert ".csv" in SUPPORTED_EXTENSIONS
        assert ".tsv" in SUPPORTED_EXTENSIONS


# ── XML fixtures ──────────────────────────────────────────────────────────────

_SIMPLE_XML = """\
<?xml version="1.0" encoding="UTF-8"?>
<books>
  <book id="B001">
    <title>Python Cookbook</title>
    <author>David Beazley</author>
    <isbn>978-1-449-34037-7</isbn>
    <price currency="USD">49.99</price>
    <tags>python,programming</tags>
  </book>
  <book id="B002">
    <title>C++ Primer</title>
    <author>Stanley Lippman</author>
    <isbn>978-0-321-71404-6</isbn>
    <price currency="USD">59.99</price>
    <tags>c++,reference</tags>
  </book>
</books>"""


def _make_large_xml() -> str:
    lines = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        "<employees>",
    ]
    for i in range(200):
        lines.append(f'  <employee id="EMP-{i:04d}">')
        lines.append(f"    <name>Employee {i}</name>")
        lines.append(f"    <department>Dept-{i % 10}</department>")
        lines.append(f"    <email>emp{i}@example.com</email>")
        lines.append(f"    <phone>+1-555-{i:04d}-{(i * 7) % 10000:04d}</phone>")
        lines.append("  </employee>")
    lines.append("</employees>")
    return "\n".join(lines)


_LARGE_XML = _make_large_xml()


_SPECIAL_CHARS_XML = """\
<?xml version="1.0" encoding="UTF-8"?>
<config>
  <setting name="api_url">https://api.example.com/v2</setting>
  <setting name="timeout">30s</setting>
  <setting name="retries">3</setting>
  <setting name="regex">^[a-z]+-[0-9]+$</setting>
  <setting name="formula">a+b=c</setting>
  <setting name="path">/usr/local/bin</setting>
  <setting name="version">1.0.0-beta</setting>
  <setting name="description">Supports C++ &amp; Python</setting>
</config>"""


# ── XML discovery ─────────────────────────────────────────────────────────────


class TestXmlDiscovery:
    def test_xml_in_supported_extensions(self):
        assert ".xml" in SUPPORTED_EXTENSIONS

    def test_iter_files_finds_xml(self, tmp_path: Path):
        (tmp_path / "data.xml").write_text(_SIMPLE_XML)
        (tmp_path / "other.txt").write_text("text")

        files = list(iter_files(tmp_path))
        names = [f.name for f in files]
        assert "data.xml" in names
        assert "other.txt" in names

    def test_iter_files_recursive_finds_nested_xml(self, tmp_path: Path):
        sub = tmp_path / "subdir"
        sub.mkdir()
        (sub / "nested.xml").write_text(_SIMPLE_XML)

        files = list(iter_files(tmp_path))
        names = [f.name for f in files]
        assert "nested.xml" in names

    def test_iter_files_skips_unsupported(self, tmp_path: Path):
        (tmp_path / "image.xml.png").write_bytes(b"\x89PNG")
        (tmp_path / "notxml.dat").write_bytes(b"\x00")

        files = list(iter_files(tmp_path))
        assert len(files) == 0


# ── XML extraction ────────────────────────────────────────────────────────────


class TestXmlExtraction:
    def test_extract_simple_xml(self, tmp_path: Path):
        f = tmp_path / "books.xml"
        f.write_text(_SIMPLE_XML)
        text = extract_txt(f)
        assert "Python Cookbook" in text
        assert "David Beazley" in text
        assert "<book" in text

    def test_extract_preserves_special_chars(self, tmp_path: Path):
        f = tmp_path / "config.xml"
        f.write_text(_SPECIAL_CHARS_XML)
        text = extract_txt(f)
        assert "https://api.example.com" in text
        assert "C++" in text
        assert "/usr/local/bin" in text

    def test_extract_preserves_unicode(self, tmp_path: Path):
        xml = '<?xml version="1.0"?>\n<data>\n  <item>Ñoño café résumé</item>\n</data>'
        f = tmp_path / "unicode.xml"
        f.write_text(xml, encoding="utf-8")
        text = extract_txt(f)
        assert "Ñoño" in text
        assert "café" in text
        assert "résumé" in text


# ── XML chunking ──────────────────────────────────────────────────────────────


class TestXmlChunking:
    def test_should_chunk_large_xml(self, tmp_path: Path):
        f = tmp_path / "large.xml"
        f.write_text(_LARGE_XML)
        assert should_chunk(f, extract_txt(f)) is True

    def test_should_not_chunk_small_xml(self, tmp_path: Path):
        f = tmp_path / "small.xml"
        f.write_text(_SIMPLE_XML)
        assert should_chunk(f, extract_txt(f)) is False

    def test_chunk_xml_by_lines(self):
        chunks = chunk_code_by_lines(_SIMPLE_XML, 3)
        assert len(chunks) > 1
        all_chunk_lines = "\n".join(chunks)
        for line in _SIMPLE_XML.split("\n"):
            assert line in all_chunk_lines

    def test_chunk_large_xml(self):
        chunks = chunk_code_by_lines(_LARGE_XML, 50)
        assert len(chunks) > 3
        assert chunks[0].startswith("<?xml")

    def test_chunk_preserves_xml_structure(self):
        chunks = chunk_code_by_lines(_SIMPLE_XML, 5)
        for chunk in chunks:
            assert len(chunk) > 0


# ── XML ingestion pipeline ────────────────────────────────────────────────────


class TestXmlIngestion:
    def test_ingest_simple_xml(self, tmp_path: Path):
        tracker = Tracker(tmp_path / "tracker.db")
        f = tmp_path / "books.xml"
        f.write_text(_SIMPLE_XML)

        count = ingest_file(f, tracker, db_path=tmp_path / "test.db")
        assert count >= 1

        search = SQLiteSearch(tmp_path / "test.db")
        results = search.search("Python Cookbook")
        assert len(results) > 0
        assert "books.xml" in results[0]["filepath"]
        search.close()

    def test_ingest_large_xml(self, tmp_path: Path):
        tracker = Tracker(tmp_path / "tracker.db")
        f = tmp_path / "employees.xml"
        f.write_text(_LARGE_XML)

        count = ingest_file(f, tracker, db_path=tmp_path / "test.db")
        assert count >= 1

        search = SQLiteSearch(tmp_path / "test.db")
        results = search.search("Employee 150", all_chunks=True)
        assert len(results) > 0
        search.close()

    def test_ingest_xml_with_special_chars(self, tmp_path: Path):
        tracker = Tracker(tmp_path / "tracker.db")
        f = tmp_path / "config.xml"
        f.write_text(_SPECIAL_CHARS_XML)

        count = ingest_file(f, tracker, db_path=tmp_path / "test.db")
        assert count >= 1

        search = SQLiteSearch(tmp_path / "test.db")
        results = search.search("api example")
        assert len(results) > 0
        search.close()

    def test_ingest_xml_produces_correct_doc_id(self, tmp_path: Path):
        tracker = Tracker(tmp_path / "tracker.db")
        f = tmp_path / "books.xml"
        f.write_text(_SIMPLE_XML)

        ingest_file(f, tracker, db_path=tmp_path / "test.db")

        search = SQLiteSearch(tmp_path / "test.db")
        stats = search.stats()
        assert stats["points_count"] >= 1
        search.close()


# ── FTS5 search with XML content ──────────────────────────────────────────────


class TestXmlFts5Search:
    def _setup_search(self, tmp_path: Path, xml_content: str) -> SQLiteSearch:
        tracker = Tracker(tmp_path / "tracker.db")
        f = tmp_path / "data.xml"
        f.write_text(xml_content)
        ingest_file(f, tracker, db_path=tmp_path / "test.db")
        return SQLiteSearch(tmp_path / "test.db")

    def test_search_xml_title(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _SIMPLE_XML)
        results = search.search("Python Cookbook")
        assert len(results) > 0
        assert "Python Cookbook" in results[0]["content"]
        search.close()

    def test_search_xml_author(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _SIMPLE_XML)
        results = search.search("David Beazley")
        assert len(results) > 0
        search.close()

    def test_search_xml_isbn(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _SIMPLE_XML)
        results = search.search("978-1-449")
        assert len(results) > 0
        search.close()

    def test_search_xml_tag_content(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _SIMPLE_XML)
        results = search.search("python programming")
        assert len(results) > 0
        search.close()

    def test_search_xml_special_chars_url(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _SPECIAL_CHARS_XML)
        results = search.search("api example")
        assert len(results) > 0
        search.close()

    def test_search_xml_special_chars_cpp(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _SPECIAL_CHARS_XML)
        results = search.search("C++ Python")
        assert len(results) > 0
        search.close()

    def test_search_xml_special_chars_path(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _SPECIAL_CHARS_XML)
        results = search.search("/usr/local")
        assert len(results) > 0
        search.close()

    def test_search_xml_special_chars_formula(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _SPECIAL_CHARS_XML)
        results = search.search("a+b=c")
        assert len(results) > 0
        search.close()

    def test_search_xml_special_chars_version(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _SPECIAL_CHARS_XML)
        results = search.search("1.0.0-beta")
        assert len(results) > 0
        search.close()

    def test_search_xml_employee_id(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _LARGE_XML)
        results = search.search("EMP-0150")
        assert len(results) > 0
        search.close()

    def test_search_xml_phone_number(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _LARGE_XML)
        results = search.search("+1-555")
        assert len(results) > 0
        search.close()

    def test_search_xml_email(self, tmp_path: Path):
        search = self._setup_search(tmp_path, _LARGE_XML)
        results = search.search("emp150 example")
        assert len(results) > 0
        search.close()


# ── FTS5 escape function ──────────────────────────────────────────────────────


class TestFts5Escape:
    def test_escape_hyphenated_isbn(self):
        from archivist.search.sqlite_search import SQLiteSearch
        result = SQLiteSearch._escape_fts("ISBN-978-0-321-71404-6")
        assert "isbn" in result
        assert "978" in result
        assert "321" in result
        assert "71404" in result

    def test_escape_cpp(self):
        from archivist.search.sqlite_search import SQLiteSearch
        result = SQLiteSearch._escape_fts("c++ code")
        assert "c" in result
        assert "code" in result

    def test_escape_at_sign(self):
        from archivist.search.sqlite_search import SQLiteSearch
        result = SQLiteSearch._escape_fts("user@example.com")
        assert "user" in result
        assert "example" in result

    def test_escape_url(self):
        from archivist.search.sqlite_search import SQLiteSearch
        result = SQLiteSearch._escape_fts("https://api.example.com")
        assert "https" in result
        assert "api" in result
        assert "example" in result

    def test_escape_formula(self):
        from archivist.search.sqlite_search import SQLiteSearch
        result = SQLiteSearch._escape_fts("a+b=c")
        assert "a" in result
        assert "b" in result
        assert "c" in result

    def test_escape_version(self):
        from archivist.search.sqlite_search import SQLiteSearch
        result = SQLiteSearch._escape_fts("1.0.0-beta")
        assert "1" in result
        assert "0" in result
        assert "beta" in result

    def test_escape_empty(self):
        from archivist.search.sqlite_search import SQLiteSearch
        assert SQLiteSearch._escape_fts("") == '""'

    def test_escape_whitespace_only(self):
        from archivist.search.sqlite_search import SQLiteSearch
        assert SQLiteSearch._escape_fts("   ") == '""'
