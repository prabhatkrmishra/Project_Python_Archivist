"""Tests for text extraction, normalization, and chunking.

Covers: extract_txt, extract_pdf, extract_docx, normalize_text, chunk_text,
chunk_pdf_by_page, chunk_docx_by_section, should_chunk.
"""
from __future__ import annotations

from pathlib import Path

import pytest

from archivist.ingestion.extractors import (
    chunk_docx_by_section,
    chunk_pdf_by_page,
    chunk_text,
    extract_docx,
    extract_pdf,
    extract_txt,
    normalize_text,
    should_chunk,
)
from docx import Document



@pytest.fixture
def txt_file(tmp_path: Path) -> Path:
    f = tmp_path / "sample.txt"
    f.write_text("Hello world\n\n\tTabbed content  \n  spaced  ")
    return f


_MINIMAL_PDF = (
    b"%PDF-1.4\n"
    b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    b"2 0 obj\n<< /Type /Pages /Kids [3 0 R 4 0 R 5 0 R] /Count 3 >>\nendobj\n"
    b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] >>\nendobj\n"
    b"4 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] >>\nendobj\n"
    b"5 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] >>\nendobj\n"
    b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n"
    b"0000000115 00000 n \n0000000216 00000 n \n0000000317 00000 n \n"
    b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n400\n%%EOF\n"
)


@pytest.fixture
def pdf_file(tmp_path: Path) -> Path:
    f = tmp_path / "sample.pdf"
    f.write_bytes(_MINIMAL_PDF)
    return f


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    f = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Introduction", 0)
    doc.add_paragraph("First paragraph.")
    doc.add_heading("Methods", 1)
    doc.add_paragraph("Second paragraph under heading.")
    doc.add_paragraph("Third paragraph.")
    doc.save(str(f))
    return f


@pytest.fixture
def large_pdf_file(tmp_path: Path) -> Path:
    """Create a file >10 MB with .pdf extension to trigger size-based chunking."""
    f = tmp_path / "large.pdf"
    f.write_bytes(b"x" * (11 * 1024 * 1024))  # 11 MB
    return f


# ── normalize_text ──────────────────────────────────────────────────────────


class TestNormalizeText:
    def test_lowercases(self):
        assert normalize_text("HELLO World") == "hello world"

    def test_collapses_whitespace(self):
        assert normalize_text("a  b\n\n\tc") == "a b c"

    def test_strips_control_chars(self):
        assert "\x00" not in normalize_text("hello\x00world")

    def test_empty_string(self):
        assert normalize_text("") == ""

    def test_strip_surrounding_whitespace(self):
        assert normalize_text("  spaced  ") == "spaced"


# ── extract_txt ─────────────────────────────────────────────────────────────


class TestExtractTxt:
    def test_utf8(self, tmp_path: Path):
        f = tmp_path / "u.txt"
        f.write_text("héllo wörld", encoding="utf-8")
        assert extract_txt(f) == "héllo wörld"

    def test_fallback_encoding(self, tmp_path: Path):
        f = tmp_path / "latin1.txt"
        f.write_bytes("café".encode("latin-1"))
        result = extract_txt(f)
        assert "caf" in result

    def test_returns_string(self, txt_file: Path):
        result = extract_txt(txt_file)
        assert isinstance(result, str)


# ── extract_pdf ──────────────────────────────────────────────────────────────


class TestExtractPdf:
    def test_returns_string(self, pdf_file: Path):
        result = extract_pdf(pdf_file)
        assert isinstance(result, str)

    def test_three_blank_pages(self, pdf_file: Path):
        result = extract_pdf(pdf_file)
        assert result.count("\n") >= 2

    def test_nonexistent_file_raises(self, tmp_path: Path):
        with pytest.raises(Exception):
            extract_pdf(tmp_path / "does_not_exist.pdf")


# ── extract_docx ─────────────────────────────────────────────────────────────


class TestExtractDocx:
    def test_returns_string(self, docx_file: Path):
        result = extract_docx(docx_file)
        assert isinstance(result, str)

    def test_contains_heading_text(self, docx_file: Path):
        result = extract_docx(docx_file)
        assert "introduction" in result.lower()

    def test_contains_paragraph_text(self, docx_file: Path):
        result = extract_docx(docx_file)
        assert "first paragraph" in result.lower()


# ── chunk_pdf_by_page ────────────────────────────────────────────────────────


class TestChunkPdfByPage:
    def test_returns_list(self, pdf_file: Path):
        chunks = chunk_pdf_by_page(pdf_file)
        assert isinstance(chunks, list)

    def test_three_pages_produces_three_chunks(self, pdf_file: Path):
        chunks = chunk_pdf_by_page(pdf_file)
        assert len(chunks) == 3

    def test_each_chunk_is_string(self, pdf_file: Path):
        chunks = chunk_pdf_by_page(pdf_file)
        assert all(isinstance(c, str) for c in chunks)


# ── chunk_docx_by_section ────────────────────────────────────────────────────


class TestChunkDocxBySection:
    def test_returns_list(self, docx_file: Path):
        chunks = chunk_docx_by_section(docx_file)
        assert isinstance(chunks, list)

    def test_splits_on_heading(self, docx_file: Path):
        chunks = chunk_docx_by_section(docx_file)
        assert len(chunks) >= 2

    def test_first_chunk_contains_introduction(self, docx_file: Path):
        chunks = chunk_docx_by_section(docx_file)
        assert any("introduction" in c.lower() for c in chunks)

    def test_second_chunk_contains_methods(self, docx_file: Path):
        chunks = chunk_docx_by_section(docx_file)
        assert any("methods" in c.lower() for c in chunks)


# ── chunk_text dispatcher ────────────────────────────────────────────────────


class TestChunkText:
    def test_txt_returns_single_chunk(self, txt_file: Path):
        text = extract_txt(txt_file)
        chunks = chunk_text(txt_file, text)
        assert len(chunks) == 1
        assert chunks[0] == text

    def test_pdf_splits_by_page(self, pdf_file: Path):
        text = extract_pdf(pdf_file)
        chunks = chunk_text(pdf_file, text)
        assert len(chunks) == 3

    def test_docx_splits_by_section(self, docx_file: Path):
        text = extract_docx(docx_file)
        chunks = chunk_text(docx_file, text)
        assert len(chunks) >= 2


# ── should_chunk ─────────────────────────────────────────────────────────────


class TestShouldChunk:
    def test_small_txt_returns_false(self, txt_file: Path):
        assert not should_chunk(txt_file, extract_txt(txt_file))

    def test_small_pdf_returns_false(self, pdf_file: Path):
        assert not should_chunk(pdf_file, extract_pdf(pdf_file))

    def test_large_pdf_by_size_returns_true(self, large_pdf_file: Path):
        assert should_chunk(large_pdf_file, "")
